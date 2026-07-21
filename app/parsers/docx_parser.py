"""Extracts raw text from .docx SOW files using python-docx."""

from pathlib import Path

from docx import Document

from app.utils.logger import get_logger

logger = get_logger(__name__)


def extract_text_from_docx(file_path: str | Path) -> str:
    """Extract all paragraph and table text from a .docx file, in order.

    Raises:
        ValueError: if the file cannot be opened as a valid docx.
    """
    try:
        document = Document(file_path)
    except Exception as exc:
        logger.error(f"Failed to open DOCX file '{file_path}': {exc}")
        raise ValueError(f"Could not read DOCX file: {exc}") from exc

    chunks: list[str] = []

    for paragraph in document.paragraphs:
        if paragraph.text.strip():
            chunks.append(paragraph.text)

    for table in document.tables:
        for row in table.rows:
            row_text = " | ".join(cell.text.strip() for cell in row.cells)
            if row_text.strip(" |"):
                chunks.append(row_text)

    text = "\n".join(chunks)
    logger.info(f"Extracted {len(text)} characters from DOCX '{file_path}'")
    return text
