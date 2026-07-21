"""
BRD Document Generator.

WHY: The AI agent outputs the BRD as Markdown-flavored text ("# Title",
"## Section", "- bullet"). This module is the ONLY place responsible for
turning that text into a polished .docx file, so formatting rules (fonts,
heading styles, page numbers) live in one place instead of being duplicated
wherever export happens.

This is intentionally a light markdown->docx converter (headings + bullets +
tables via "|"-separated rows), not a full markdown engine — that's enough
for the structure our own prompt templates produce.
"""

import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt

from app.utils.logger import get_logger

logger = get_logger(__name__)


def _add_page_numbers(document: Document) -> None:
    """Add 'Page X of Y' to the footer of every page."""
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run = paragraph.add_run("Page ")

    def _field(field_code: str):
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = field_code
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)

    _field("PAGE")
    paragraph.add_run(" of ")
    _field("NUMPAGES")


def _looks_like_table_row(line: str) -> bool:
    return "|" in line.strip()


def _add_table_from_rows(document: Document, rows: list[str]) -> None:
    parsed_rows = [
        [cell.strip() for cell in row.strip().strip("|").split("|")]
        for row in rows
        if row.strip()
    ]
    if not parsed_rows:
        return

    col_count = max(len(r) for r in parsed_rows)
    table = document.add_table(rows=0, cols=col_count)
    table.style = "Light Grid Accent 1"

    for row_values in parsed_rows:
        row_cells = table.add_row().cells
        for idx in range(col_count):
            row_cells[idx].text = row_values[idx] if idx < len(row_values) else ""


def generate_brd_docx(brd_markdown: str, output_path: str | Path) -> Path:
    """Convert a markdown-flavored BRD string into a formatted .docx file.

    Supports:
        # / ## / ### headings
        - / * bullet points
        1. numbered list items
        | table | rows |  (simple pipe-delimited tables)
        plain paragraphs
    """
    document = Document()

    # Base document style
    style = document.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    lines = brd_markdown.splitlines()
    i = 0
    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        if not stripped:
            i += 1
            continue

        # Collect contiguous table rows together
        if _looks_like_table_row(stripped) and not stripped.startswith(("#", "-", "*")):
            table_rows = []
            while i < len(lines) and _looks_like_table_row(lines[i].strip()) and lines[i].strip():
                table_rows.append(lines[i])
                i += 1
            _add_table_from_rows(document, table_rows)
            continue

        heading_match = re.match(r"^(#{1,3})\s+(.*)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            document.add_heading(text, level=level)
            i += 1
            continue

        bullet_match = re.match(r"^[-*]\s+(.*)", stripped)
        if bullet_match:
            document.add_paragraph(bullet_match.group(1), style="List Bullet")
            i += 1
            continue

        numbered_match = re.match(r"^\d+[.)]\s+(.*)", stripped)
        if numbered_match:
            document.add_paragraph(numbered_match.group(1), style="List Number")
            i += 1
            continue

        bold_meta_match = re.match(r"^\*\*(.+?):\*\*\s*(.*)", stripped)
        if bold_meta_match:
            p = document.add_paragraph()
            run = p.add_run(f"{bold_meta_match.group(1)}: ")
            run.bold = True
            p.add_run(bold_meta_match.group(2))
            i += 1
            continue

        document.add_paragraph(stripped)
        i += 1

    _add_page_numbers(document)

    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)

    logger.info(f"Generated BRD .docx at '{output_path}'")
    return output_path
